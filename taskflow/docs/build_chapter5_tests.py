# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = "docs/Chuong_5_Kiem_thu_TaskFlow.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
HEADER_FILL = "F2F4F7"
BORDER = "B7C3D0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_grid(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_table(table, widths, header_rows=1):
    set_table_grid(table, widths)
    set_table_borders(table)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)
            if r_idx < header_rows:
                set_cell_shading(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text).italic = True
    return p


def add_table(doc, caption, headers, rows, widths, header_rows=1):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, value in enumerate(headers):
        table.cell(0, idx).text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table, widths, header_rows=header_rows)
    doc.add_paragraph()
    return table


def add_test_table(doc, caption, rows):
    headers = ["TC", "Mục đích kiểm thử", "Các bước thực hiện", "Kết quả mong đợi", "Kết quả thực hiện", "Pass/Fail"]
    widths = [650, 2050, 2750, 2450, 900, 560]
    return add_table(doc, caption, headers, rows, widths)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    styles["Caption"].font.name = "Calibri"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.color.rgb = RGBColor(85, 85, 85)
    styles["Caption"].paragraph_format.space_before = Pt(4)
    styles["Caption"].paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "TaskFlow - Chương 5. Kiểm thử"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def build_doc():
    doc = Document()
    configure_document(doc)

    doc.add_heading("CHƯƠNG 5. KIỂM THỬ", level=1)
    doc.add_paragraph(
        "Trong chương này nhóm em xây dựng các ca kiểm thử cho hệ thống TaskFlow đã triển khai. "
        "Nội dung tập trung vào kiểm thử chức năng của các màn hình chính, các thao tác liên quan đến Supabase "
        "và một số luồng có thể kiểm thử bằng Postman hoặc công cụ kiểm thử giao diện."
    )
    doc.add_paragraph(
        "Do phạm vi hiện tại là xây dựng ca kiểm thử, cột kết quả thực hiện được để là “Chưa thực hiện”. "
        "Khi chạy kiểm thử thực tế, nhóm có thể cập nhật lại cột này thành kết quả thực tế và đánh dấu Pass/Fail."
    )

    doc.add_heading("5.1 Kiểm thử chức năng", level=2)

    doc.add_heading("5.1.1 Kiểm thử chức năng đăng nhập", level=3)
    doc.add_paragraph(
        "Khi đăng nhập vào hệ thống, người dùng nhập email và mật khẩu đã được đăng ký. "
        "Bộ kiểm thử dựa trên hai trường email và mật khẩu với ba trạng thái: hợp lệ, không hợp lệ và bỏ trống."
    )
    add_table(
        doc,
        "Bảng 5.1 Bảng quyết định cho kiểm thử đăng nhập",
        ["Điều kiện", "1", "2", "3", "4", "5", "6", "7", "8", "9"],
        [
            ["Email", "hợp lệ", "hợp lệ", "hợp lệ", "không hợp lệ", "không hợp lệ", "không hợp lệ", "Trống", "Trống", "Trống"],
            ["Mật khẩu", "hợp lệ", "không hợp lệ", "Trống", "hợp lệ", "không hợp lệ", "Trống", "hợp lệ", "không hợp lệ", "Trống"],
            ["Kết quả", "", "", "", "", "", "", "", "", ""],
            ["Login Pass", "x", "", "", "", "", "", "", "", ""],
            ["Login Fail", "", "x", "x", "x", "x", "x", "x", "x", "x"],
        ],
        [1300, 895, 895, 895, 895, 895, 895, 895, 895, 895],
    )
    add_table(
        doc,
        "Bảng 5.2 Bảng quyết định cho kiểm thử đăng nhập sau rút gọn",
        ["Điều kiện", "1", "2", "3", "4"],
        [
            ["Email", "hợp lệ", "hợp lệ", "không hợp lệ", "Trống"],
            ["Mật khẩu", "hợp lệ", "-", "-", "-"],
            ["Kết quả", "", "", "", ""],
            ["Login Pass", "x", "", "", ""],
            ["Login Fail", "", "x", "x", "x"],
        ],
        [2200, 1790, 1790, 1790, 1790],
    )
    add_test_table(
        doc,
        "Bảng 5.3 Bảng ca kiểm thử chức năng đăng nhập",
        [
            ["TC01", "Đăng nhập đúng thông tin", "1. Mở trang /login\n2. Nhập email hợp lệ\n3. Nhập mật khẩu đúng\n4. Bấm đăng nhập", "Đăng nhập thành công và chuyển đến dashboard", "Chưa thực hiện", ""],
            ["TC02", "Email đúng nhưng mật khẩu sai", "1. Nhập email tồn tại\n2. Nhập mật khẩu sai\n3. Bấm đăng nhập", "Hệ thống không cho đăng nhập và hiển thị thông báo lỗi", "Chưa thực hiện", ""],
            ["TC03", "Email không tồn tại", "1. Nhập email chưa đăng ký\n2. Nhập mật khẩu bất kỳ\n3. Bấm đăng nhập", "Hệ thống không cho đăng nhập", "Chưa thực hiện", ""],
            ["TC04", "Bỏ trống email hoặc mật khẩu", "1. Bỏ trống một trong hai trường\n2. Bấm đăng nhập", "Hệ thống yêu cầu nhập đầy đủ thông tin", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.2 Kiểm thử chức năng đăng ký và hồ sơ cá nhân", level=3)
    doc.add_paragraph("Áp dụng phương pháp phân vùng tương đương cho các trường họ tên, email, mật khẩu, số điện thoại và ngày sinh.")
    add_table(
        doc,
        "Bảng 5.4 Bảng phân vùng tương đương chức năng đăng ký và hồ sơ cá nhân",
        ["STT", "Đầu vào", "Lớp hợp lệ", "Ký hiệu", "Vùng không hợp lệ", "Ký hiệu"],
        [
            ["1", "Họ tên", "Không bỏ trống", "H1", "Bỏ trống", "K1"],
            ["2", "Email", "Đúng định dạng email", "H2", "Sai định dạng hoặc đã tồn tại", "K2"],
            ["3", "Mật khẩu", "Từ 6 ký tự trở lên", "H3", "Dưới 6 ký tự hoặc bỏ trống", "K3"],
            ["4", "Số điện thoại", "10 chữ số", "H4", "Chứa chữ, ký tự đặc biệt hoặc sai độ dài", "K4"],
            ["5", "Ngày sinh", "Định dạng ngày/tháng/năm", "H5", "Sai định dạng hoặc ngày không tồn tại", "K5"],
        ],
        [650, 1500, 2300, 850, 3000, 1060],
    )
    add_test_table(
        doc,
        "Bảng 5.5 Bảng ca kiểm thử chức năng đăng ký và hồ sơ cá nhân",
        [
            ["TC01", "Đăng ký tài khoản hợp lệ", "Nhập đầy đủ họ tên, email mới, mật khẩu hợp lệ rồi bấm đăng ký", "Tạo tài khoản thành công và có hồ sơ người dùng", "Chưa thực hiện", ""],
            ["TC02", "Email sai định dạng", "Nhập email không có ký tự @ rồi bấm đăng ký", "Hiển thị lỗi email không hợp lệ", "Chưa thực hiện", ""],
            ["TC03", "Email đã tồn tại", "Nhập email đã có trong hệ thống rồi bấm đăng ký", "Không tạo tài khoản trùng", "Chưa thực hiện", ""],
            ["TC04", "Mật khẩu quá ngắn", "Nhập mật khẩu dưới 6 ký tự rồi bấm đăng ký", "Hiển thị lỗi mật khẩu không hợp lệ", "Chưa thực hiện", ""],
            ["TC05", "Cập nhật hồ sơ hợp lệ", "Mở trang hồ sơ, sửa họ tên, số điện thoại, ngày sinh dạng dd/mm/yyyy rồi lưu", "Thông tin được cập nhật và hiển thị lại đúng định dạng ngày/tháng/năm", "Chưa thực hiện", ""],
            ["TC06", "Ngày sinh sai định dạng", "Nhập ngày sinh dạng mm/dd/yyyy hoặc ngày không tồn tại rồi lưu", "Không lưu dữ liệu sai và hiển thị thông báo phù hợp", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.3 Kiểm thử chức năng quản lý dự án", level=3)
    add_test_table(
        doc,
        "Bảng 5.6 Bảng ca kiểm thử chức năng quản lý dự án",
        [
            ["TC01", "Tạo dự án mới hợp lệ", "1. Mở trang tạo dự án\n2. Nhập tên và mô tả\n3. Bấm tạo", "Dự án được tạo, người tạo là owner/member của dự án", "Chưa thực hiện", ""],
            ["TC02", "Tạo dự án thiếu tên", "Bỏ trống tên dự án rồi bấm tạo", "Không tạo dự án và hiển thị lỗi bắt buộc nhập tên", "Chưa thực hiện", ""],
            ["TC03", "Cập nhật dự án", "Mở chi tiết dự án, sửa tên hoặc mô tả rồi lưu", "Thông tin dự án được cập nhật trong danh sách và trang chi tiết", "Chưa thực hiện", ""],
            ["TC04", "Xóa dự án", "Chọn dự án, thực hiện xóa và xác nhận", "Dự án bị xóa, task và thành viên liên quan được xử lý theo nghiệp vụ", "Chưa thực hiện", ""],
            ["TC05", "Người không thuộc dự án truy cập chi tiết", "Dùng tài khoản không phải thành viên mở đường dẫn /projects/[id]", "Không cho xem dữ liệu dự án hoặc chuyển về trang phù hợp", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.4 Kiểm thử chức năng quản lý thành viên dự án", level=3)
    add_test_table(
        doc,
        "Bảng 5.7 Bảng ca kiểm thử chức năng quản lý thành viên dự án",
        [
            ["TC01", "Mời thành viên bằng email hợp lệ", "Nhập email người dùng đã tồn tại rồi bấm mời", "Thành viên được thêm vào dự án", "Chưa thực hiện", ""],
            ["TC02", "Mời email không tồn tại", "Nhập email không có trong hệ thống rồi bấm mời", "Hiển thị thông báo không tìm thấy người dùng", "Chưa thực hiện", ""],
            ["TC03", "Mời trùng thành viên", "Mời lại người dùng đã thuộc dự án", "Không tạo bản ghi trùng thành viên", "Chưa thực hiện", ""],
            ["TC04", "Tìm kiếm người dùng", "Nhập từ khóa tên hoặc email vào ô tìm kiếm thành viên", "Danh sách gợi ý hiển thị đúng người dùng phù hợp", "Chưa thực hiện", ""],
            ["TC05", "Phân quyền truy cập theo vai trò", "Dùng tài khoản member thực hiện thao tác chỉ owner/admin được phép làm", "Hệ thống chặn thao tác không đúng quyền", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.5 Kiểm thử chức năng quản lý công việc", level=3)
    add_test_table(
        doc,
        "Bảng 5.8 Bảng ca kiểm thử chức năng quản lý công việc",
        [
            ["TC01", "Tạo task hợp lệ", "Trong dự án, nhập tiêu đề, mô tả, hạn hoàn thành rồi bấm tạo", "Task xuất hiện trong cột trạng thái mặc định", "Chưa thực hiện", ""],
            ["TC02", "Tạo task thiếu tiêu đề", "Bỏ trống tiêu đề task rồi bấm tạo", "Không tạo task và thông báo lỗi", "Chưa thực hiện", ""],
            ["TC03", "Sửa task", "Mở task, cập nhật tiêu đề hoặc mô tả rồi lưu", "Thông tin task được cập nhật", "Chưa thực hiện", ""],
            ["TC04", "Kéo thả đổi trạng thái", "Kéo task từ To do sang In progress hoặc Done", "Trạng thái và vị trí task được lưu lại", "Chưa thực hiện", ""],
            ["TC05", "Hoàn thành task", "Bấm hoàn thành task trong danh sách công việc", "Task chuyển sang trạng thái hoàn thành", "Chưa thực hiện", ""],
            ["TC06", "Xóa task", "Chọn xóa task và xác nhận", "Task bị xóa khỏi danh sách", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.6 Kiểm thử chức năng phân công và công việc của tôi", level=3)
    add_test_table(
        doc,
        "Bảng 5.9 Bảng ca kiểm thử chức năng phân công task",
        [
            ["TC01", "Gán người thực hiện", "Mở task, chọn thành viên dự án để gán", "Người được gán hiển thị trong task", "Chưa thực hiện", ""],
            ["TC02", "Gán người không thuộc dự án", "Thử gán user không nằm trong danh sách thành viên dự án", "Hệ thống không cho gán", "Chưa thực hiện", ""],
            ["TC03", "Gỡ người thực hiện", "Mở task đã có assignee, chọn gỡ thành viên", "Thành viên bị gỡ khỏi task", "Chưa thực hiện", ""],
            ["TC04", "Xem trang My Tasks", "Đăng nhập bằng user được gán task rồi mở /my-tasks", "Chỉ hiển thị các task được gán cho user hiện tại", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.7 Kiểm thử chức năng bình luận trong task", level=3)
    add_test_table(
        doc,
        "Bảng 5.10 Bảng ca kiểm thử chức năng bình luận task",
        [
            ["TC01", "Thêm bình luận hợp lệ", "Mở task, nhập nội dung bình luận rồi gửi", "Bình luận xuất hiện kèm người gửi và thời gian", "Chưa thực hiện", ""],
            ["TC02", "Gửi bình luận rỗng", "Để trống nội dung rồi bấm gửi", "Không tạo bình luận rỗng", "Chưa thực hiện", ""],
            ["TC03", "Sửa bình luận của mình", "Chọn bình luận do mình gửi, sửa nội dung rồi lưu", "Nội dung bình luận được cập nhật", "Chưa thực hiện", ""],
            ["TC04", "Xóa bình luận của mình", "Chọn bình luận do mình gửi và xóa", "Bình luận bị xóa khỏi task", "Chưa thực hiện", ""],
            ["TC05", "Sửa/xóa bình luận của người khác", "Dùng tài khoản khác thử sửa hoặc xóa bình luận", "Hệ thống không cho thao tác sai quyền", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.8 Kiểm thử chức năng tin nhắn", level=3)
    add_test_table(
        doc,
        "Bảng 5.11 Bảng ca kiểm thử chức năng tin nhắn",
        [
            ["TC01", "Tải danh sách người dùng có thể nhắn", "Mở trang Messages sau khi đăng nhập", "Danh sách người dùng hiển thị đúng, có avatar nếu đã cấu hình", "Chưa thực hiện", ""],
            ["TC02", "Gửi tin nhắn văn bản", "Chọn một người nhận, nhập nội dung rồi gửi", "Tin nhắn hiển thị ở cuộc hội thoại và lưu vào Supabase", "Chưa thực hiện", ""],
            ["TC03", "Không gửi tin nhắn rỗng", "Chọn người nhận, để trống nội dung rồi gửi", "Không tạo tin nhắn rỗng", "Chưa thực hiện", ""],
            ["TC04", "Nhận tin nhắn realtime", "Mở hai tài khoản ở hai trình duyệt, tài khoản A gửi tin cho B", "Tài khoản B nhận tin mới mà không cần tải lại trang", "Chưa thực hiện", ""],
            ["TC05", "Tìm kiếm trong tin nhắn", "Nhập từ khóa vào ô tìm kiếm tin nhắn", "Danh sách tin nhắn hoặc hội thoại được lọc theo từ khóa", "Chưa thực hiện", ""],
            ["TC06", "Upload file/ảnh trong tin nhắn", "Chọn file hợp lệ rồi gửi trong cuộc hội thoại", "File được upload và liên kết hiển thị trong tin nhắn", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.9 Kiểm thử chức năng thông báo", level=3)
    add_test_table(
        doc,
        "Bảng 5.12 Bảng ca kiểm thử chức năng thông báo",
        [
            ["TC01", "Tạo thông báo khi có sự kiện", "Thực hiện thao tác tạo task hoặc gán task cho user", "Thông báo được tạo cho người liên quan", "Chưa thực hiện", ""],
            ["TC02", "Xem danh sách thông báo", "Mở trang Notifications", "Danh sách thông báo của user hiện tại được hiển thị", "Chưa thực hiện", ""],
            ["TC03", "Không xem thông báo của user khác", "Dùng user khác truy cập dữ liệu thông báo", "Không hiển thị thông báo không thuộc user hiện tại", "Chưa thực hiện", ""],
            ["TC04", "Cập nhật trạng thái đã đọc", "Mở hoặc đánh dấu thông báo đã đọc", "Thông báo đổi trạng thái từ chưa đọc sang đã đọc", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.10 Kiểm thử chức năng quản trị tài khoản", level=3)
    add_test_table(
        doc,
        "Bảng 5.13 Bảng ca kiểm thử chức năng quản trị tài khoản",
        [
            ["TC01", "Admin xem danh sách người dùng", "Đăng nhập bằng admin, mở trang Admin", "Danh sách người dùng hiển thị đầy đủ thông tin cần quản lý", "Chưa thực hiện", ""],
            ["TC02", "Cập nhật vai trò người dùng", "Chọn một user, đổi role và lưu", "Vai trò mới được cập nhật", "Chưa thực hiện", ""],
            ["TC03", "Khóa/mở khóa tài khoản", "Admin đổi trạng thái disabled của user", "User bị khóa không thể sử dụng hệ thống hoặc được mở lại đúng trạng thái", "Chưa thực hiện", ""],
            ["TC04", "Đổi mật khẩu người dùng", "Admin nhập mật khẩu mới hợp lệ cho user", "Mật khẩu được cập nhật", "Chưa thực hiện", ""],
            ["TC05", "User thường truy cập trang admin", "Đăng nhập bằng user thường và mở /admin", "Hệ thống chặn truy cập hoặc chuyển hướng", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.11 Kiểm thử chức năng dashboard và thống kê", level=3)
    add_test_table(
        doc,
        "Bảng 5.14 Bảng ca kiểm thử chức năng dashboard và thống kê",
        [
            ["TC01", "Hiển thị tổng số dự án", "Tạo thêm một dự án rồi mở dashboard", "Số lượng dự án tăng tương ứng", "Chưa thực hiện", ""],
            ["TC02", "Hiển thị tổng số task", "Tạo thêm task trong dự án rồi mở dashboard", "Số lượng task tăng tương ứng", "Chưa thực hiện", ""],
            ["TC03", "Thống kê task theo trạng thái", "Đổi trạng thái task rồi mở analytics", "Biểu đồ/thống kê cập nhật đúng theo trạng thái mới", "Chưa thực hiện", ""],
            ["TC04", "Không hiển thị dữ liệu ngoài phạm vi quyền", "User không thuộc dự án mở dashboard", "Không thống kê dữ liệu dự án không được phép xem", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.12 Kiểm thử chức năng upload hình ảnh và media", level=3)
    add_test_table(
        doc,
        "Bảng 5.15 Bảng ca kiểm thử chức năng upload hình ảnh và media",
        [
            ["TC01", "Upload avatar hợp lệ", "Chọn file ảnh hợp lệ tại trang hồ sơ rồi lưu", "Ảnh được upload và avatar mới hiển thị", "Chưa thực hiện", ""],
            ["TC02", "Upload file không phải ảnh", "Chọn file sai định dạng rồi upload", "Hệ thống từ chối hoặc hiển thị lỗi định dạng", "Chưa thực hiện", ""],
            ["TC03", "Upload file quá lớn", "Chọn file vượt giới hạn cho phép", "Hệ thống không upload và thông báo lỗi", "Chưa thực hiện", ""],
            ["TC04", "Xóa ảnh đã upload", "Thực hiện thao tác xóa ảnh media", "Ảnh bị xóa khỏi dịch vụ lưu trữ nếu nghiệp vụ cho phép", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.1.13 Kiểm thử API danh sách dự án bằng Postman", level=3)
    doc.add_paragraph(
        "Với các API hoặc endpoint có thể gọi trực tiếp, nhóm sử dụng Postman để gửi request, kiểm tra mã trạng thái, "
        "schema dữ liệu trả về và phân quyền truy cập."
    )
    add_test_table(
        doc,
        "Bảng 5.16 Bảng ca kiểm thử API dự án bằng Postman",
        [
            ["TC01", "Lấy danh sách dự án thành công", "Gửi GET /api/projects với môi trường local đang chạy", "Trả về HTTP 200 và danh sách dự án đúng cấu trúc", "Chưa thực hiện", ""],
            ["TC02", "Kiểm tra schema dữ liệu", "Gửi request lấy danh sách dự án, kiểm tra các trường id, name, description", "Response có đủ trường cần thiết, kiểu dữ liệu phù hợp", "Chưa thực hiện", ""],
            ["TC03", "Request sai phương thức", "Gửi POST hoặc PUT tới endpoint chỉ hỗ trợ GET", "Trả về mã lỗi phù hợp, không làm thay đổi dữ liệu", "Chưa thực hiện", ""],
            ["TC04", "Kiểm tra thời gian phản hồi", "Gửi request lặp lại bằng Postman Runner", "Thời gian phản hồi nằm trong ngưỡng chấp nhận của nhóm", "Chưa thực hiện", ""],
        ],
    )

    doc.add_heading("5.2 Kiểm thử hiệu năng", level=2)
    doc.add_heading("5.2.1 Giới thiệu công cụ kiểm thử", level=3)
    doc.add_paragraph(
        "Đối với kiểm thử hiệu năng, nhóm có thể sử dụng JMeter để mô phỏng nhiều người dùng đồng thời. "
        "Postman Runner có thể dùng để kiểm tra nhanh API, nhưng JMeter phù hợp hơn khi cần đo tải, throughput, "
        "thời gian phản hồi trung bình và tỷ lệ lỗi."
    )
    doc.add_heading("5.2.2 Kiểm thử hiệu năng hệ thống TaskFlow", level=3)
    doc.add_paragraph(
        "Các luồng được đề xuất kiểm thử gồm đăng nhập, tải dashboard, mở danh sách dự án, tạo task và gửi tin nhắn. "
        "Môi trường kiểm thử là ứng dụng chạy local hoặc staging, sử dụng dữ liệu mẫu đã chuẩn bị trước."
    )
    add_table(
        doc,
        "Bảng 5.17 Bảng kịch bản kiểm thử hiệu năng",
        ["STT", "Luồng kiểm thử", "Request/Thao tác chính", "Chỉ số cần đo", "Ngưỡng mong muốn"],
        [
            ["1", "Đăng nhập", "POST tới Supabase Auth hoặc thao tác login UI", "Average response time, error rate", "Tỷ lệ lỗi dưới 1%, phản hồi trung bình dưới 2 giây"],
            ["2", "Dashboard", "Tải trang /dashboard", "Load time, response time", "Trang tải ổn định dưới 3 giây với tải vừa"],
            ["3", "Danh sách dự án", "GET /api/projects hoặc truy vấn danh sách dự án", "Throughput, response time", "Không lỗi dữ liệu, phản hồi dưới 2 giây"],
            ["4", "Tạo task", "Gửi thao tác tạo task trong một dự án", "Response time, error rate", "Không tạo task trùng/lỗi khi nhiều request đồng thời"],
            ["5", "Tin nhắn", "Gửi tin nhắn giữa hai tài khoản", "Response time, realtime delay", "Tin nhắn lưu thành công, realtime delay dưới 3 giây"],
        ],
        [650, 1650, 2650, 2150, 2260],
    )
    add_table(
        doc,
        "Bảng 5.18 Bảng mức tải đề xuất cho kiểm thử hiệu năng",
        ["Mức tải", "Số người dùng đồng thời", "Thời gian chạy", "Kết quả cần ghi nhận"],
        [
            ["Nhẹ", "10 user", "5 phút", "Thời gian phản hồi trung bình, lỗi phát sinh"],
            ["Trung bình", "50 user", "10 phút", "Throughput, CPU/RAM nếu có, tỷ lệ lỗi"],
            ["Cao", "100 user", "15 phút", "Điểm bắt đầu chậm, tỷ lệ lỗi, khả năng phục hồi"],
        ],
        [1200, 2200, 2200, 3760],
    )
    doc.add_paragraph(
        "Sau khi chạy kiểm thử hiệu năng, nhóm điền số liệu thực tế gồm tổng user, tổng thời gian chạy, "
        "thời gian phản hồi trung bình, thời gian phản hồi nhỏ nhất/lớn nhất, tỷ lệ request lỗi và throughput."
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
